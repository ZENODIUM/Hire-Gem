import os
import json
from flask import Flask, render_template, request, jsonify, send_from_directory
from werkzeug.utils import secure_filename
import requests
from bs4 import BeautifulSoup
import google.generativeai as genai
from google.generativeai import types
from daytona import Daytona, DaytonaConfig
import PyPDF2
from docx import Document
import re
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

try:
    from firecrawl import Firecrawl
    FIRECRAWL_AVAILABLE = True
except ImportError:
    FIRECRAWL_AVAILABLE = False
    print("Firecrawl not available. Install with: pip install firecrawl-py")

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['SECRET_KEY'] = 'your-secret-key-here'

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Data storage directory for saved profiles
DATA_STORAGE_DIR = 'profile_data'
os.makedirs(DATA_STORAGE_DIR, exist_ok=True)

# Configure Gemini AI
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY environment variable is required. Please set it in your .env file.")

genai.configure(api_key=GEMINI_API_KEY)
# Try user-specified model, fallback to available models
try:
    model = genai.GenerativeModel('gemini-2.5-flash')
except:
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
    except:
        model = genai.GenerativeModel('gemini-pro')

# Initialize Daytona
DAYTONA_API_KEY = os.getenv('DAYTONA_API_KEY')
if DAYTONA_API_KEY:
    daytona_config = DaytonaConfig(api_key=DAYTONA_API_KEY)
    daytona = Daytona(daytona_config)
else:
    daytona = None
    print("Warning: DAYTONA_API_KEY not set. Daytona features will be disabled.")

# Initialize Firecrawl
FIRECRAWL_API_KEY = os.getenv('FIRECRAWL_API_KEY')
if FIRECRAWL_AVAILABLE and FIRECRAWL_API_KEY:
    try:
        firecrawl = Firecrawl(api_key=FIRECRAWL_API_KEY)
    except Exception as e:
        print(f"Error initializing Firecrawl: {e}")
        firecrawl = None
else:
    firecrawl = None
    if not FIRECRAWL_API_KEY:
        print("Warning: FIRECRAWL_API_KEY not set. Firecrawl features will be disabled.")

ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return text

def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print(f"Error reading TXT: {e}")
        return ""

def extract_links_from_pdf(file_path):
    """Extract URLs from PDF file (both text and hyperlinks)"""
    links = set()
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                # Extract text and find URLs in text
                text = page.extract_text()
                url_pattern = re.compile(r'https?://[^\s<>"{}|\\^`\[\]]+[^\s<>"{}|\\^`\[\].,;:!?]')
                text_urls = url_pattern.findall(text)
                links.update(text_urls)
                
                # Extract hyperlinks from annotations
                if '/Annots' in page:
                    annotations = page['/Annots']
                    for annotation in annotations:
                        obj = annotation.get_object()
                        if '/A' in obj and '/URI' in obj['/A']:
                            uri = obj['/A']['/URI']
                            if isinstance(uri, str) and (uri.startswith('http://') or uri.startswith('https://')):
                                links.add(uri)
    except Exception as e:
        print(f"Error extracting links from PDF: {e}")
    return list(links)

def extract_links_from_docx(file_path):
    """Extract URLs from DOCX file (both text and hyperlinks)"""
    links = set()
    try:
        doc = Document(file_path)
        
        # Extract URLs from text
        for paragraph in doc.paragraphs:
            text = paragraph.text
            url_pattern = re.compile(r'https?://[^\s<>"{}|\\^`\[\]]+[^\s<>"{}|\\^`\[\].,;:!?]')
            text_urls = url_pattern.findall(text)
            links.update(text_urls)
            
            # Extract hyperlinks from runs
            for run in paragraph.runs:
                if run.hyperlink and run.hyperlink.address:
                    links.add(run.hyperlink.address)
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        url_pattern = re.compile(r'https?://[^\s<>"{}|\\^`\[\]]+[^\s<>"{}|\\^`\[\].,;:!?]')
                        text_urls = url_pattern.findall(text)
                        links.update(text_urls)
    except Exception as e:
        print(f"Error extracting links from DOCX: {e}")
    return list(links)

def extract_links_from_txt(file_path):
    """Extract URLs from TXT file"""
    links = set()
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            url_pattern = re.compile(r'https?://[^\s<>"{}|\\^`\[\]]+[^\s<>"{}|\\^`\[\].,;:!?]')
            text_urls = url_pattern.findall(content)
            links.update(text_urls)
    except Exception as e:
        print(f"Error extracting links from TXT: {e}")
    return list(links)

def extract_links_from_file(file_path, file_ext):
    """Extract all URLs from a file based on its extension"""
    if file_ext == 'pdf':
        return extract_links_from_pdf(file_path)
    elif file_ext == 'docx':
        return extract_links_from_docx(file_path)
    else:
        return extract_links_from_txt(file_path)

def scrape_github(username):
    """Scrape GitHub profile information using GitHub API"""
    try:
        # Use Daytona sandbox for secure execution if available
        if daytona:
            try:
                # Create scraping code to run in sandbox
                scraping_code = f"""
import requests
import json
import base64

api_url = "https://api.github.com/users/{username}"
response = requests.get(api_url, timeout=10)

if response.status_code == 200:
    data = response.json()
    repos_url = f"https://api.github.com/users/{username}/repos?sort=updated&per_page=10"
    repos_response = requests.get(repos_url, timeout=10)
    repos = repos_response.json() if repos_response.status_code == 200 else []
    
    detailed_repos = []
    for repo in repos[:10]:
        repo_name = repo.get('name', '')
        repo_full_name = repo.get('full_name', '')
        
        repo_info = {{
            'name': repo_name,
            'full_name': repo_full_name,
            'description': repo.get('description', ''),
            'language': repo.get('language', ''),
            'stars': repo.get('stargazers_count', 0),
            'forks': repo.get('forks_count', 0),
            'url': repo.get('html_url', ''),
            'topics': repo.get('topics', []),
            'detailed_description': ''
        }}
        
        # Try to get README
        try:
            readme_url = f"https://api.github.com/repos/{{repo_full_name}}/readme"
            readme_response = requests.get(readme_url, timeout=10)
            if readme_response.status_code == 200:
                readme_data = readme_response.json()
                readme_content = base64.b64decode(readme_data.get('content', '')).decode('utf-8')
                repo_info['detailed_description'] = readme_content[:1000] + '...' if len(readme_content) > 1000 else readme_content
        except:
            pass
        
        detailed_repos.append(repo_info)
    
    profile_info = {{
        'name': data.get('name', ''),
        'bio': data.get('bio', ''),
        'location': data.get('location', ''),
        'company': data.get('company', ''),
        'blog': data.get('blog', ''),
        'public_repos': data.get('public_repos', 0),
        'followers': data.get('followers', 0),
        'following': data.get('following', 0),
        'repositories': detailed_repos
    }}
    print(json.dumps(profile_info))
else:
    print(json.dumps({{'error': f'GitHub API returned status {{response.status_code}}'}}))
"""
                sandbox = daytona.create()
                result = sandbox.process.code_run(scraping_code)
                sandbox.delete()
                
                if result.exit_code == 0:
                    import json
                    return json.loads(result.result)
            except Exception as e:
                print(f"Daytona execution failed, falling back to direct API: {e}")
        
        # Fallback to direct API call
        api_url = f"https://api.github.com/users/{username}"
        response = requests.get(api_url, timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            
            # Get repositories
            repos_url = f"https://api.github.com/users/{username}/repos?sort=updated&per_page=10"
            repos_response = requests.get(repos_url, timeout=10)
            repos = repos_response.json() if repos_response.status_code == 200 else []
            
            # Get detailed information for each repository
            detailed_repos = []
            for repo in repos[:10]:
                repo_name = repo.get('name', '')
                repo_full_name = repo.get('full_name', '')
                
                repo_info = {
                    'name': repo_name,
                    'full_name': repo_full_name,
                    'description': repo.get('description', ''),
                    'language': repo.get('language', ''),
                    'stars': repo.get('stargazers_count', 0),
                    'forks': repo.get('forks_count', 0),
                    'url': repo.get('html_url', ''),
                    'created_at': repo.get('created_at', ''),
                    'updated_at': repo.get('updated_at', ''),
                    'topics': repo.get('topics', []),
                    'detailed_description': ''
                }
                
                # Try to get README content for detailed description
                try:
                    readme_url = f"https://api.github.com/repos/{repo_full_name}/readme"
                    readme_response = requests.get(readme_url, timeout=10)
                    if readme_response.status_code == 200:
                        readme_data = readme_response.json()
                        import base64
                        readme_content = base64.b64decode(readme_data.get('content', '')).decode('utf-8')
                        # Get first 1000 characters of README
                        repo_info['detailed_description'] = readme_content[:1000] + '...' if len(readme_content) > 1000 else readme_content
                except Exception as e:
                    # If README fetch fails, try to get more details from repo API
                    try:
                        repo_detail_url = f"https://api.github.com/repos/{repo_full_name}"
                        repo_detail_response = requests.get(repo_detail_url, timeout=10)
                        if repo_detail_response.status_code == 200:
                            repo_detail = repo_detail_response.json()
                            # Use homepage or more detailed description
                            if repo_detail.get('homepage'):
                                repo_info['homepage'] = repo_detail.get('homepage')
                            if repo_detail.get('description') and not repo_info['description']:
                                repo_info['description'] = repo_detail.get('description', '')
                    except:
                        pass
                
                detailed_repos.append(repo_info)
            
            profile_info = {
                'name': data.get('name', ''),
                'bio': data.get('bio', ''),
                'location': data.get('location', ''),
                'company': data.get('company', ''),
                'blog': data.get('blog', ''),
                'public_repos': data.get('public_repos', 0),
                'followers': data.get('followers', 0),
                'following': data.get('following', 0),
                'repositories': detailed_repos
            }
            return profile_info
        else:
            return {'error': f'GitHub API returned status {response.status_code}'}
    except Exception as e:
        return {'error': f'Error scraping GitHub: {str(e)}'}

def scrape_linkedin(profile_url):
    """Scrape LinkedIn profile information"""
    try:
        # LinkedIn requires authentication for API access
        # For scraping, we'll use a simplified approach
        # Note: LinkedIn has strict anti-scraping measures
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(profile_url, headers=headers, timeout=10)
        
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract basic info (LinkedIn structure may vary)
            profile_info = {
                'name': '',
                'headline': '',
                'location': '',
                'about': '',
                'experience': [],
                'education': []
            }
            
            # Try to extract name
            name_elem = soup.find('h1', class_='text-heading-xlarge')
            if name_elem:
                profile_info['name'] = name_elem.get_text(strip=True)
            
            # Try to extract headline
            headline_elem = soup.find('div', class_='text-body-medium')
            if headline_elem:
                profile_info['headline'] = headline_elem.get_text(strip=True)
            
            # Note: LinkedIn's HTML structure changes frequently
            # For production, consider using LinkedIn API or specialized tools
            return profile_info
        else:
            return {'error': f'LinkedIn returned status {response.status_code}'}
    except Exception as e:
        return {'error': f'Error scraping LinkedIn: {str(e)}'}

def scrape_devpost_project(project_url):
    """Scrape individual DevPost project page for detailed information"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        }
        response = requests.get(project_url, headers=headers, timeout=15)
        
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            
            project_details = {
                'url': project_url,
                'youtube_links': [],
                'video_links': [],
                'full_description': '',
                'built_with': [],
                'team_members': [],
                'screenshots': []
            }
            
            # Extract YouTube links - multiple methods
            # Method 1: Find iframe embeds
            iframes = soup.find_all('iframe')
            for iframe in iframes:
                src = iframe.get('src', '')
                if 'youtube.com' in src or 'youtu.be' in src:
                    # Extract YouTube video ID
                    youtube_match = re.search(r'(?:youtube\.com\/embed\/|youtu\.be\/)([a-zA-Z0-9_-]+)', src)
                    if youtube_match:
                        video_id = youtube_match.group(1)
                        project_details['youtube_links'].append(f'https://www.youtube.com/watch?v={video_id}')
            
            # Method 2: Find YouTube links in text/links
            all_links = soup.find_all('a', href=True)
            for link in all_links:
                href = link.get('href', '')
                if 'youtube.com' in href or 'youtu.be' in href:
                    # Normalize YouTube URL
                    if 'youtu.be' in href:
                        video_id = href.split('/')[-1].split('?')[0]
                        normalized = f'https://www.youtube.com/watch?v={video_id}'
                    elif 'watch?v=' in href:
                        normalized = href.split('&')[0]  # Remove extra parameters
                    else:
                        normalized = href
                    
                    if normalized not in project_details['youtube_links']:
                        project_details['youtube_links'].append(normalized)
            
            # Method 3: Find in text content (sometimes URLs are in text)
            text_content = soup.get_text()
            youtube_pattern = re.compile(r'(?:https?://)?(?:www\.)?(?:youtube\.com/watch\?v=|youtu\.be/)([a-zA-Z0-9_-]+)')
            matches = youtube_pattern.findall(text_content)
            for video_id in matches:
                normalized = f'https://www.youtube.com/watch?v={video_id}'
                if normalized not in project_details['youtube_links']:
                    project_details['youtube_links'].append(normalized)
            
            # Extract full description
            description_elem = soup.find('div', class_=re.compile(r'description|about|overview', re.I))
            if description_elem:
                project_details['full_description'] = description_elem.get_text(strip=True)[:2000]
            
            # Extract built with technologies
            built_with_elem = soup.find(['div', 'section'], class_=re.compile(r'built-with|technologies|tech-stack', re.I))
            if built_with_elem:
                tech_tags = built_with_elem.find_all(['span', 'a', 'li'], class_=re.compile(r'tag|tech|technology', re.I))
                for tag in tech_tags:
                    tech = tag.get_text(strip=True)
                    if tech and tech not in project_details['built_with']:
                        project_details['built_with'].append(tech)
            
            # Extract team members
            team_elem = soup.find(['div', 'section'], class_=re.compile(r'team|contributors|authors', re.I))
            if team_elem:
                member_links = team_elem.find_all('a', href=re.compile(r'/users/'))
                for member_link in member_links:
                    member_name = member_link.get_text(strip=True)
                    if member_name and member_name not in project_details['team_members']:
                        project_details['team_members'].append(member_name)
            
            # Extract screenshots/gallery images
            gallery = soup.find(['div', 'section'], class_=re.compile(r'gallery|screenshots|images', re.I))
            if gallery:
                imgs = gallery.find_all('img')
                for img in imgs:
                    img_src = img.get('src') or img.get('data-src')
                    if img_src and 'screenshot' in img_src.lower() or 'gallery' in img_src.lower():
                        if img_src.startswith('//'):
                            img_src = 'https:' + img_src
                        elif img_src.startswith('/'):
                            img_src = 'https://devpost.com' + img_src
                        project_details['screenshots'].append(img_src)
            
            return project_details
        else:
            return {'error': f'Failed to fetch project page: {response.status_code}'}
    except Exception as e:
        return {'error': f'Error scraping project page: {str(e)}'}

def scrape_devpost(username):
    """Scrape DevPost profile information using web scraping and AI parsing"""
    try:
        url = f"https://devpost.com/{username}"
        
        # Try Firecrawl first for better content extraction
        firecrawl_result = scrape_with_firecrawl(url)
        use_firecrawl = False
        
        if firecrawl_result and firecrawl_result.get('markdown'):
            # Use Firecrawl content but still need to parse HTML for images
            use_firecrawl = True
            visible_text_firecrawl = firecrawl_result.get('markdown', '')[:8000]
        
        # Always get HTML for image extraction
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=15)
        
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract project images - more comprehensive approach
            project_images = []
            project_image_map = {}  # Map project names/links to images
            project_urls_list = []  # Store URLs for later use
            
            # Method 1: Find all links to software/projects
            project_links = soup.find_all('a', href=re.compile(r'/software/'))
            for link in project_links:
                img = link.find('img')
                if img:
                    img_url = img.get('src') or img.get('data-src') or img.get('data-original') or img.get('data-lazy-src')
                    if img_url:
                        # Normalize URL
                        if img_url.startswith('//'):
                            img_url = 'https:' + img_url
                        elif img_url.startswith('/'):
                            img_url = 'https://devpost.com' + img_url
                        elif not img_url.startswith('http'):
                            img_url = 'https://devpost.com/' + img_url
                        
                        # Get project identifier from href
                        href = link.get('href', '')
                        project_slug = href.split('/')[-1] if href else None
                        
                        # Store full URL for later scraping
                        if href:
                            if href.startswith('/'):
                                full_url = 'https://devpost.com' + href
                            elif href.startswith('http'):
                                full_url = href
                            else:
                                full_url = 'https://devpost.com/' + href
                            if full_url not in project_urls_list:
                                project_urls_list.append(full_url)
                        
                        # Get project name from nearby text
                        project_name = None
                        name_elem = link.find(['h5', 'h4', 'h3', 'h2', 'span', 'div'], class_=re.compile(r'name|title|heading', re.I))
                        if name_elem:
                            project_name = name_elem.get_text(strip=True)
                        elif img.get('alt'):
                            project_name = img.get('alt')
                        
                        if img_url and img_url not in [img['url'] for img in project_images]:
                            project_images.append({
                                'url': img_url,
                                'project_slug': project_slug,
                                'project_name': project_name,
                                'href': href
                            })
                            
                            # Store mapping for easier lookup
                            if project_slug:
                                project_image_map[project_slug.lower()] = img_url
                            if project_name:
                                project_image_map[project_name.lower().strip()] = img_url
            
            # Method 2: Find images in project-related containers
            containers = soup.find_all(['div', 'article', 'section'], class_=re.compile(r'software|project|entry|card', re.I))
            for container in containers:
                # Check if container has a project link
                has_project_link = container.find('a', href=re.compile(r'/software/')) is not None
                imgs = container.find_all('img')
                for img in imgs:
                    img_url = img.get('src') or img.get('data-src') or img.get('data-original') or img.get('data-lazy-src')
                    if img_url and (has_project_link or 'software' in img_url.lower() or 'project' in img_url.lower()):
                        if img_url.startswith('//'):
                            img_url = 'https:' + img_url
                        elif img_url.startswith('/'):
                            img_url = 'https://devpost.com' + img_url
                        elif not img_url.startswith('http'):
                            img_url = 'https://devpost.com/' + img_url
                        
                        if img_url not in [img['url'] for img in project_images]:
                            # Try to find associated project link
                            link = container.find('a', href=re.compile(r'/software/'))
                            project_slug = None
                            if link:
                                href = link.get('href', '')
                                project_slug = href.split('/')[-1] if href else None
                            
                            project_images.append({
                                'url': img_url,
                                'project_slug': project_slug,
                                'project_name': img.get('alt', ''),
                                'href': link.get('href', '') if link else ''
                            })
                            
                            if project_slug:
                                project_image_map[project_slug.lower()] = img_url
            
            # Remove script and style elements
            for script in soup(["script", "style", "meta", "link"]):
                script.decompose()
            
            # Get all text content
            text_content = soup.get_text(separator='\n', strip=True)
            
            # Extract visible text (limit to avoid token limits)
            visible_text = '\n'.join([line.strip() for line in text_content.split('\n') if line.strip()])[:8000]
            
            # Prepare image info for AI
            image_info = f"Found {len(project_images)} project images. Image URLs: " + ", ".join([img['url'] for img in project_images[:10]])
            
            # Use AI to parse and structure the DevPost data
            structure_prompt = f"""Extract and structure the following DevPost profile information into JSON format.

Raw HTML text content from DevPost profile page:
{visible_text}

{image_info}

Please extract and return a JSON object with the following structure:
{{
    "name": "Full name of the person",
    "location": "Location if available",
    "bio": "Bio or description if available",
    "skills": ["skill1", "skill2", ...],
    "interests": ["interest1", "interest2", ...],
    "stats": {{
        "projects": number,
        "hackathons": number,
        "achievements": number,
        "followers": number,
        "following": number
    }},
    "projects": [
        {{
            "name": "Project name",
            "description": "Project description",
            "technologies": ["tech1", "tech2"],
            "awards": ["award info if any"]
        }}
    ],
    "linkedin_url": "LinkedIn URL if present"
}}

Return ONLY valid JSON, no additional text or markdown formatting."""

            try:
                ai_response = model.generate_content(structure_prompt)
                ai_text = ai_response.text.strip()
                
                # Clean the response (remove markdown code blocks if present)
                if ai_text.startswith('```'):
                    ai_text = ai_text.split('```')[1]
                    if ai_text.startswith('json'):
                        ai_text = ai_text[4:]
                ai_text = ai_text.strip()
                
                # Parse JSON
                structured_data = json.loads(ai_text)
                
                # Add images to projects
                if 'projects' in structured_data:
                    used_images = set()
                    for idx, project in enumerate(structured_data['projects']):
                        project_name = project.get('name', '').lower().strip()
                        matched = False
                        
                        # Try multiple matching strategies
                        # Strategy 1: Match by project name
                        if project_name:
                            for img_data in project_images:
                                img_name = (img_data.get('project_name') or '').lower().strip()
                                img_slug = (img_data.get('project_slug') or '').lower().strip()
                                img_url = img_data.get('url', '')
                                
                                if img_url and img_url not in used_images:
                                    # Check if names match
                                    if (project_name and img_name and 
                                        (project_name in img_name or img_name in project_name or
                                         project_name[:15] in img_name or img_name[:15] in project_name)):
                                        project['image_url'] = img_url
                                        used_images.add(img_url)
                                        matched = True
                                        break
                                    
                                    # Check if slug matches project name
                                    if project_name and img_slug and project_name.replace(' ', '-') in img_slug:
                                        project['image_url'] = img_url
                                        used_images.add(img_url)
                                        matched = True
                                        break
                        
                        # Strategy 2: Match by index (assign images in order)
                        if not matched and idx < len(project_images):
                            img_data = project_images[idx]
                            img_url = img_data.get('url', '')
                            if img_url and img_url not in used_images:
                                project['image_url'] = img_url
                                used_images.add(img_url)
                                matched = True
                        
                        # Strategy 3: Assign any unused image
                        if not matched:
                            for img_data in project_images:
                                img_url = img_data.get('url', '')
                                if img_url and img_url not in used_images:
                                    project['image_url'] = img_url
                                    used_images.add(img_url)
                                    break
                    
                    # Ensure all projects get images if available (fallback by index)
                    if project_images:
                        for idx, project in enumerate(structured_data['projects']):
                            if 'image_url' not in project and idx < len(project_images):
                                project['image_url'] = project_images[idx]['url']
                    
                    # Debug: Add image count info (remove before returning in production)
                    # structured_data['_debug'] = {
                    #     'images_found': len(project_images),
                    #     'projects_count': len(structured_data.get('projects', []))
                    # }
                
                # Also add raw image list for debugging
                if project_images:
                    structured_data['_images'] = [img['url'] for img in project_images[:20]]
                
                # Scrape detailed information for top 3 projects
                if 'projects' in structured_data and len(structured_data['projects']) > 0:
                    # Use the project URLs we collected earlier
                    project_urls = project_urls_list[:3]  # Top 3 projects
                    
                    # Scrape each project page
                    for idx, project in enumerate(structured_data['projects'][:3]):
                        if idx < len(project_urls):
                            project_url = project_urls[idx]
                            try:
                                project_details = scrape_devpost_project(project_url)
                                # Merge details into project
                                if 'youtube_links' in project_details and project_details['youtube_links']:
                                    project['youtube_links'] = project_details['youtube_links']
                                if 'full_description' in project_details and project_details['full_description']:
                                    project['full_description'] = project_details['full_description']
                                if 'built_with' in project_details and project_details['built_with']:
                                    project['built_with'] = project_details['built_with']
                                if 'team_members' in project_details and project_details['team_members']:
                                    project['team_members'] = project_details['team_members']
                                if 'screenshots' in project_details and project_details['screenshots']:
                                    project['screenshots'] = project_details['screenshots']
                                project['project_url'] = project_url
                            except Exception as e:
                                print(f"Error scraping project details for {project_url}: {e}")
                        else:
                            # If we don't have URL, try to construct from project name
                            project_slug = project.get('name', '').lower().replace(' ', '-')
                            if project_slug:
                                project['project_url'] = f'https://devpost.com/software/{project_slug}'
                
                return structured_data
            except json.JSONDecodeError as e:
                # Fallback: try to extract basic info manually
                profile_info = {
                    'name': '',
                    'bio': '',
                    'skills': [],
                    'projects': [],
                    'raw_text': visible_text[:2000]  # Include raw text as fallback
                }
                
                # Try to find name in h1 tags
                h1_tags = soup.find_all('h1')
                for h1 in h1_tags:
                    text = h1.get_text(strip=True)
                    if text and len(text) < 100:
                        profile_info['name'] = text
                        break
                
                return profile_info
        else:
            return {'error': f'DevPost returned status {response.status_code}'}
    except Exception as e:
        return {'error': f'Error scraping DevPost: {str(e)}'}

def scrape_portfolio(url):
    """Scrape general portfolio/website content using web scraping and AI parsing"""
    try:
        # Try Firecrawl first
        firecrawl_result = scrape_with_firecrawl(url)
        
        if firecrawl_result:
            # Use Firecrawl's markdown content
            visible_text = firecrawl_result.get('markdown', '') or firecrawl_result.get('content', '')
            # Limit text for AI processing
            visible_text = visible_text[:10000] if len(visible_text) > 10000 else visible_text
        else:
            # Fallback to direct scraping
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
            }
            response = requests.get(url, headers=headers, timeout=15)
            
            if response.status_code != 200:
                return {'error': f'Website returned status {response.status_code}'}
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "meta", "link", "nav", "footer"]):
                script.decompose()
            
            # Get main content
            main_content = soup.find('main') or soup.find('body') or soup
            text_content = main_content.get_text(separator='\n', strip=True)
            
            # Extract visible text (limit to avoid token limits)
            visible_text = '\n'.join([line.strip() for line in text_content.split('\n') if line.strip() and len(line.strip()) > 3])[:10000]
        
        # Use AI to parse and structure the portfolio data (common for both Firecrawl and direct scraping)
        structure_prompt = f"""Analyze the following website/portfolio content and extract professional information into JSON format.

Raw text content from the website:
{visible_text}

Please extract and return a JSON object with the following structure:
{{
    "name": "Name if available",
    "title": "Professional title or role",
    "bio": "Bio, about, or introduction text",
    "skills": ["skill1", "skill2", ...],
    "technologies": ["tech1", "tech2", ...],
    "projects": [
        {{
            "name": "Project name",
            "description": "Project description",
            "technologies": ["tech1", "tech2"]
        }}
    ],
    "experience": [
        {{
            "role": "Job title",
            "company": "Company name",
            "description": "Description"
        }}
    ],
    "education": ["education item 1", "education item 2"],
    "contact": {{
        "email": "email if found",
        "linkedin": "LinkedIn URL if found",
        "github": "GitHub URL if found"
    }},
    "summary": "Brief summary of the portfolio content"
}}

Return ONLY valid JSON, no additional text or markdown formatting."""

        try:
            ai_response = model.generate_content(structure_prompt)
            ai_text = ai_response.text.strip()
            
            # Clean the response (remove markdown code blocks if present)
            if ai_text.startswith('```'):
                ai_text = ai_text.split('```')[1]
                if ai_text.startswith('json'):
                    ai_text = ai_text[4:]
            ai_text = ai_text.strip()
            
            # Parse JSON
            structured_data = json.loads(ai_text)
            structured_data['source_url'] = url
            return structured_data
        except json.JSONDecodeError as e:
            # Fallback: return raw text
            return {
                'source_url': url,
                'raw_content': visible_text[:3000],
                'error': 'Could not parse structured data, showing raw content'
            }
    except Exception as e:
        return {'error': f'Error scraping portfolio: {str(e)}'}

def scrape_kaggle(username):
    """Scrape Kaggle profile information from main page, code, and datasets pages"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        }
        
        kaggle_data = {
            'username': username,
            'profile': {},
            'code': [],
            'datasets': []
        }
        
        # Scrape main profile page
        profile_url = f"https://www.kaggle.com/{username}"
        try:
            response = requests.get(profile_url, headers=headers, timeout=15)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style", "meta", "link"]):
                    script.decompose()
                
                # Extract profile information
                text_content = soup.get_text(separator='\n', strip=True)
                visible_text = '\n'.join([line.strip() for line in text_content.split('\n') if line.strip()])[:5000]
                
                # Use AI to extract profile info
                profile_prompt = f"""Extract Kaggle profile information from the following text:

{visible_text}

Return a JSON object with:
{{
    "name": "Full name",
    "bio": "Bio or description",
    "location": "Location if available",
    "occupation": "Occupation if available",
    "organization": "Organization if available",
    "competitions": {{
        "tier": "Competition tier (Novice, Contributor, Expert, Master, Grandmaster)",
        "medals": {{
            "gold": number,
            "silver": number,
            "bronze": number
        }},
        "total": number
    }},
    "datasets": {{
        "tier": "Dataset tier",
        "total": number
    }},
    "notebooks": {{
        "tier": "Notebook tier",
        "total": number
    }},
    "discussion": {{
        "tier": "Discussion tier",
        "total": number
    }},
    "followers": number,
    "following": number
}}

Return ONLY valid JSON, no markdown formatting."""
                
                try:
                    ai_response = model.generate_content(profile_prompt)
                    ai_text = ai_response.text.strip()
                    if ai_text.startswith('```'):
                        ai_text = ai_text.split('```')[1]
                        if ai_text.startswith('json'):
                            ai_text = ai_text[4:]
                    ai_text = ai_text.strip()
                    kaggle_data['profile'] = json.loads(ai_text)
                except Exception as e:
                    print(f"Error parsing Kaggle profile: {e}")
        except Exception as e:
            print(f"Error scraping Kaggle profile: {e}")
        
        # Scrape code/notebooks page
        code_url = f"https://www.kaggle.com/{username}/code"
        try:
            response = requests.get(code_url, headers=headers, timeout=15)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                
                for script in soup(["script", "style", "meta", "link"]):
                    script.decompose()
                
                text_content = soup.get_text(separator='\n', strip=True)
                visible_text = '\n'.join([line.strip() for line in text_content.split('\n') if line.strip()])[:5000]
                
                # Extract notebook/code information
                code_prompt = f"""Extract Kaggle notebooks/code information from:

{visible_text}

Return a JSON array of notebooks:
[
    {{
        "title": "Notebook title",
        "description": "Description",
        "language": "Python or R",
        "votes": number,
        "views": number,
        "last_run": "Last run date if available"
    }}
]

Return ONLY valid JSON array, no markdown formatting."""
                
                try:
                    ai_response = model.generate_content(code_prompt)
                    ai_text = ai_response.text.strip()
                    if ai_text.startswith('```'):
                        ai_text = ai_text.split('```')[1]
                        if ai_text.startswith('json'):
                            ai_text = ai_text[4:]
                    ai_text = ai_text.strip()
                    kaggle_data['code'] = json.loads(ai_text)
                except Exception as e:
                    print(f"Error parsing Kaggle code: {e}")
        except Exception as e:
            print(f"Error scraping Kaggle code: {e}")
        
        # Scrape datasets page
        datasets_url = f"https://www.kaggle.com/{username}/datasets"
        try:
            response = requests.get(datasets_url, headers=headers, timeout=15)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                
                for script in soup(["script", "style", "meta", "link"]):
                    script.decompose()
                
                text_content = soup.get_text(separator='\n', strip=True)
                visible_text = '\n'.join([line.strip() for line in text_content.split('\n') if line.strip()])[:5000]
                
                # Extract datasets information
                datasets_prompt = f"""Extract Kaggle datasets information from:

{visible_text}

Return a JSON array of datasets:
[
    {{
        "title": "Dataset title",
        "description": "Description",
        "size": "Dataset size",
        "files": number,
        "downloads": number,
        "votes": number,
        "usability": number
    }}
]

Return ONLY valid JSON array, no markdown formatting."""
                
                try:
                    ai_response = model.generate_content(datasets_prompt)
                    ai_text = ai_response.text.strip()
                    if ai_text.startswith('```'):
                        ai_text = ai_text.split('```')[1]
                        if ai_text.startswith('json'):
                            ai_text = ai_text[4:]
                    ai_text = ai_text.strip()
                    kaggle_data['datasets'] = json.loads(ai_text)
                except Exception as e:
                    print(f"Error parsing Kaggle datasets: {e}")
        except Exception as e:
            print(f"Error scraping Kaggle datasets: {e}")
        
        return kaggle_data
    except Exception as e:
        return {'error': f'Error scraping Kaggle: {str(e)}'}

def scrape_with_firecrawl(url):
    """Scrape URL using Firecrawl API"""
    if not firecrawl:
        return None
    
    try:
        # Ensure URL has protocol
        if not url.startswith('http://') and not url.startswith('https://'):
            url = 'https://' + url
        
        # Use Firecrawl to scrape - correct API format
        result = firecrawl.scrape(url, formats=["markdown", "html"])
        
        # Handle response format
        if result:
            # Result is typically a dict with markdown, html, and metadata
            if isinstance(result, dict):
                return {
                    'markdown': result.get('markdown', '') or result.get('data', {}).get('markdown', ''),
                    'html': result.get('html', '') or result.get('data', {}).get('html', ''),
                    'content': result.get('content', '') or result.get('data', {}).get('content', ''),
                    'metadata': result.get('metadata', {}) or result.get('data', {}).get('metadata', {})
                }
            # If result is an object with attributes
            elif hasattr(result, 'markdown') or hasattr(result, 'content'):
                return {
                    'markdown': getattr(result, 'markdown', ''),
                    'html': getattr(result, 'html', ''),
                    'content': getattr(result, 'content', ''),
                    'metadata': getattr(result, 'metadata', {})
                }
    except Exception as e:
        print(f"Firecrawl scraping error: {e}")
    
    return None

def scrape_unknown_website(url):
    """Scrape any unknown website and use AI to identify and summarize it"""
    try:
        # Try Firecrawl first
        firecrawl_result = scrape_with_firecrawl(url)
        
        if firecrawl_result:
            # Use Firecrawl's markdown content
            visible_text = firecrawl_result.get('markdown', '') or firecrawl_result.get('content', '')
            page_title = firecrawl_result.get('metadata', {}).get('title', 'Untitled Page')
            meta_description = firecrawl_result.get('metadata', {}).get('description', '')
        else:
            # Fallback to direct scraping
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
            }
            
            # Ensure URL has protocol
            if not url.startswith('http://') and not url.startswith('https://'):
                url = 'https://' + url
            
            response = requests.get(url, headers=headers, timeout=20, allow_redirects=True)
            
            if response.status_code != 200:
                return {'error': f'Failed to fetch website: HTTP {response.status_code}'}
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove unwanted elements
            for script in soup(["script", "style", "meta", "link", "nav", "footer", "header"]):
                script.decompose()
            
            # Get main content
            main_content = soup.find('main') or soup.find('article') or soup.find('body') or soup
            text_content = main_content.get_text(separator='\n', strip=True)
            
            # Extract visible text (limit to avoid token limits but get substantial content)
            visible_text = '\n'.join([line.strip() for line in text_content.split('\n') 
                                     if line.strip() and len(line.strip()) > 3])[:15000]
            
            # Get page title
            title = soup.find('title')
            page_title = title.get_text(strip=True) if title else 'Untitled Page'
            
            # Get meta description if available
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            meta_description = meta_desc.get('content', '') if meta_desc else ''
        
        # Use AI to identify and summarize the page
        ai_prompt = f"""Analyze the following website content and provide a comprehensive summary.

Website URL: {url}
Page Title: {page_title}
Meta Description: {meta_description}

Website Content:
{visible_text}

Please analyze this content and return a JSON object with the following structure:
{{
    "page_type": "Type of page (e.g., 'Portfolio', 'Blog', 'Company Website', 'Project Page', 'Profile', 'Article', 'Documentation', 'Other')",
    "category": "Category (e.g., 'Professional Profile', 'Technical Blog', 'E-commerce', 'Educational', etc.)",
    "summary": "A comprehensive 2-3 paragraph summary of what this page is about, its purpose, and key information",
    "key_topics": ["topic1", "topic2", "topic3", ...],
    "technologies_mentioned": ["tech1", "tech2", ...],
    "skills_demonstrated": ["skill1", "skill2", ...],
    "projects_mentioned": [
        {{
            "name": "Project name",
            "description": "Brief description"
        }}
    ],
    "contact_info": {{
        "email": "email if found",
        "social_links": ["link1", "link2"]
    }},
    "main_content": "The main message or purpose of this page in 1-2 sentences",
    "professional_relevance": "How this page relates to the person's professional profile"
}}

Return ONLY valid JSON, no markdown formatting or additional text."""

        try:
            ai_response = model.generate_content(ai_prompt)
            ai_text = ai_response.text.strip()
            
            # Clean the response
            if ai_text.startswith('```'):
                ai_text = ai_text.split('```')[1]
                if ai_text.startswith('json'):
                    ai_text = ai_text[4:]
            ai_text = ai_text.strip()
            
            # Parse JSON
            analysis = json.loads(ai_text)
            
            # Add metadata
            analysis['url'] = url
            analysis['page_title'] = page_title
            analysis['meta_description'] = meta_description
            analysis['content_preview'] = visible_text[:500] + '...' if len(visible_text) > 500 else visible_text
            
            return analysis
        except json.JSONDecodeError as e:
            # Fallback: return basic analysis
            return {
                'url': url,
                'page_title': page_title,
                'meta_description': meta_description,
                'page_type': 'Unknown',
                'summary': f"This appears to be a webpage. Content preview: {visible_text[:1000]}...",
                'content_preview': visible_text[:1000],
                'error': 'Could not parse AI response, showing raw content'
            }
    except requests.exceptions.RequestException as e:
        return {'error': f'Network error: {str(e)}'}
    except Exception as e:
        return {'error': f'Error scraping website: {str(e)}'}

def extract_platform_info(url):
    """Extract platform type and identifier from URL"""
    if 'github.com' in url:
        # Extract username from GitHub URL
        match = re.search(r'github\.com/([^/]+)', url)
        if match:
            return {'platform': 'github', 'identifier': match.group(1)}
    elif 'linkedin.com' in url:
        # Extract profile identifier
        match = re.search(r'linkedin\.com/in/([^/]+)', url)
        if match:
            return {'platform': 'linkedin', 'identifier': match.group(1), 'url': url}
    elif 'devpost.com' in url:
        # Extract username from DevPost URL
        match = re.search(r'devpost\.com/([^/?]+)', url)
        if match:
            return {'platform': 'devpost', 'identifier': match.group(1)}
    elif 'kaggle.com' in url:
        # Extract username from Kaggle URL
        match = re.search(r'kaggle\.com/([^/?]+)', url)
        if match:
            return {'platform': 'kaggle', 'identifier': match.group(1)}
    else:
        # Normalize URL - ensure it has protocol
        normalized_url = url.strip()
        if not normalized_url.startswith('http://') and not normalized_url.startswith('https://'):
            if '.' in normalized_url:
                normalized_url = 'https://' + normalized_url
            else:
                # Not a valid URL format
                return None
        
        # Treat any unrecognized URL as unknown website - will use AI to identify
        return {'platform': 'unknown', 'url': normalized_url}
    
    return None

def sanitize_folder_name(name):
    """Sanitize name to be used as folder name"""
    # Remove or replace invalid characters
    sanitized = re.sub(r'[<>:"/\\|?*]', '_', name)
    sanitized = sanitized.strip()
    # Limit length
    if len(sanitized) > 100:
        sanitized = sanitized[:100]
    return sanitized or 'unknown'

def save_profile_data(person_name, cv_text, cv_file_path, analysis, scraped_data, job_description=None):
    """Save all profile data to filesystem organized by person name/ID"""
    try:
        # Sanitize person name for folder
        folder_name = sanitize_folder_name(person_name)
        person_dir = os.path.join(DATA_STORAGE_DIR, folder_name)
        os.makedirs(person_dir, exist_ok=True)
        
        # Save resume file
        if cv_file_path and os.path.exists(cv_file_path):
            resume_filename = os.path.basename(cv_file_path)
            resume_dest = os.path.join(person_dir, resume_filename)
            import shutil
            shutil.copy2(cv_file_path, resume_dest)
        
        # Save resume text
        resume_text_path = os.path.join(person_dir, 'resume_text.txt')
        with open(resume_text_path, 'w', encoding='utf-8') as f:
            f.write(cv_text)
        
        # Save analysis JSON
        analysis_path = os.path.join(person_dir, 'analysis.json')
        with open(analysis_path, 'w', encoding='utf-8') as f:
            json.dump(analysis, f, indent=2, ensure_ascii=False)
        
        # Save scraped data JSON
        scraped_data_path = os.path.join(person_dir, 'scraped_data.json')
        with open(scraped_data_path, 'w', encoding='utf-8') as f:
            json.dump(scraped_data, f, indent=2, ensure_ascii=False)
        
        # Save job description if provided
        if job_description and job_description.strip():
            job_desc_path = os.path.join(person_dir, 'job_description.txt')
            with open(job_desc_path, 'w', encoding='utf-8') as f:
                f.write(job_description)
        
        # Save metadata
        metadata = {
            'person_name': person_name,
            'folder_name': folder_name,
            'saved_at': str(os.path.getmtime(person_dir)) if os.path.exists(person_dir) else None,
            'has_resume': os.path.exists(resume_text_path),
            'has_analysis': os.path.exists(analysis_path),
            'has_scraped_data': os.path.exists(scraped_data_path),
            'has_job_description': job_description and job_description.strip() != ''
        }
        metadata_path = os.path.join(person_dir, 'metadata.json')
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2, ensure_ascii=False)
        
        return person_dir
    except Exception as e:
        print(f"Error saving profile data: {e}")
        return None

def load_profile_data(person_name):
    """Load all saved profile data for a person"""
    try:
        folder_name = sanitize_folder_name(person_name)
        person_dir = os.path.join(DATA_STORAGE_DIR, folder_name)
        
        if not os.path.exists(person_dir):
            return None
        
        data = {}
        
        # Load resume text
        resume_text_path = os.path.join(person_dir, 'resume_text.txt')
        if os.path.exists(resume_text_path):
            with open(resume_text_path, 'r', encoding='utf-8') as f:
                data['resume_text'] = f.read()
        
        # Load analysis
        analysis_path = os.path.join(person_dir, 'analysis.json')
        if os.path.exists(analysis_path):
            with open(analysis_path, 'r', encoding='utf-8') as f:
                data['analysis'] = json.load(f)
        
        # Load scraped data
        scraped_data_path = os.path.join(person_dir, 'scraped_data.json')
        if os.path.exists(scraped_data_path):
            with open(scraped_data_path, 'r', encoding='utf-8') as f:
                data['scraped_data'] = json.load(f)
        
        # Load job description
        job_desc_path = os.path.join(person_dir, 'job_description.txt')
        if os.path.exists(job_desc_path):
            with open(job_desc_path, 'r', encoding='utf-8') as f:
                data['job_description'] = f.read()
        
        # Load metadata
        metadata_path = os.path.join(person_dir, 'metadata.json')
        if os.path.exists(metadata_path):
            with open(metadata_path, 'r', encoding='utf-8') as f:
                data['metadata'] = json.load(f)
        
        return data
    except Exception as e:
        print(f"Error loading profile data: {e}")
        return None

def get_all_saved_persons():
    """Get list of all saved person names/IDs"""
    try:
        if not os.path.exists(DATA_STORAGE_DIR):
            return []
        
        persons = []
        for folder in os.listdir(DATA_STORAGE_DIR):
            folder_path = os.path.join(DATA_STORAGE_DIR, folder)
            if os.path.isdir(folder_path):
                metadata_path = os.path.join(folder_path, 'metadata.json')
                if os.path.exists(metadata_path):
                    try:
                        with open(metadata_path, 'r', encoding='utf-8') as f:
                            metadata = json.load(f)
                            persons.append({
                                'name': metadata.get('person_name', folder),
                                'folder': folder,
                                'saved_at': metadata.get('saved_at')
                            })
                    except:
                        persons.append({'name': folder, 'folder': folder})
        return persons
    except Exception as e:
        print(f"Error getting saved persons: {e}")
        return []

def generate_profile_summary(cv_text, scraped_data, job_description=None):
    """Generate a comprehensive profile summary with match analysis using Gemini AI"""
    try:
        # Prepare the prompt
        if job_description and job_description.strip():
            # Job description provided - do match analysis
            prompt = f"""Analyze the following professional profile against the job description and provide a comprehensive analysis.

JOB DESCRIPTION:
{job_description[:2000]}

CV/RESUME CONTENT:
{cv_text[:3000]}

SCRAPED PROFILE DATA (from GitHub, LinkedIn, DevPost, Kaggle, etc.):
{json.dumps(scraped_data, indent=2)[:4000]}

Please analyze and return a JSON object with the following structure:
{{
    "match_score": number (0-100, representing how well the profile matches the job),
    "summary": "A comprehensive 2-3 paragraph summary of the candidate's profile",
    "strengths": [
        "Unique strength 1 (e.g., specific projects, achievements, skills from links)",
        "Unique strength 2",
        "Unique strength 3"
    ],
    "weaknesses": [
        "Area of improvement 1 (missing skills, experience gaps, etc.)",
        "Area of improvement 2"
    ],
    "key_points": [
        "Key point 1 about the candidate",
        "Key point 2",
        "Key point 3"
    ],
    "unique_highlights": [
        "Unique thing from resume or links 1",
        "Unique thing from resume or links 2",
        "Unique thing from resume or links 3"
    ],
    "skills_match": {{
        "matched_skills": ["skill1", "skill2", ...],
        "missing_skills": ["skill1", "skill2", ...]
    }},
    "recommendations": [
        "Recommendation 1 to improve match",
        "Recommendation 2"
    ]
}}

Return ONLY valid JSON, no markdown formatting or additional text."""
        else:
            # No job description - general profile summary
            prompt = f"""Analyze the following professional profile information and create a comprehensive, well-structured summary.

CV/Resume Content:
{cv_text[:3000]}

Scraped Profile Data:
{json.dumps(scraped_data, indent=2)[:4000]}

Please analyze and return a JSON object with the following structure:
{{
    "summary": "A comprehensive 2-3 paragraph summary of the candidate's profile",
    "strengths": [
        "Unique strength 1 (e.g., specific projects, achievements, skills from links)",
        "Unique strength 2",
        "Unique strength 3"
    ],
    "weaknesses": [
        "Area of improvement 1",
        "Area of improvement 2"
    ],
    "key_points": [
        "Key point 1 about the candidate",
        "Key point 2",
        "Key point 3"
    ],
    "unique_highlights": [
        "Unique thing from resume or links 1",
        "Unique thing from resume or links 2",
        "Unique thing from resume or links 3"
    ],
    "skills": ["skill1", "skill2", ...],
    "recommendations": [
        "Recommendation 1 for professional growth",
        "Recommendation 2"
    ]
}}

Return ONLY valid JSON, no markdown formatting or additional text."""

        response = model.generate_content(prompt)
        ai_text = response.text.strip()
        
        # Clean the response
        if ai_text.startswith('```'):
            ai_text = ai_text.split('```')[1]
            if ai_text.startswith('json'):
                ai_text = ai_text[4:]
        ai_text = ai_text.strip()
        
        # Parse JSON
        try:
            analysis = json.loads(ai_text)
            return analysis
        except json.JSONDecodeError:
            # Fallback to text format
            return {
                'summary': ai_text,
                'strengths': [],
                'weaknesses': [],
                'key_points': [],
                'unique_highlights': []
            }
    except Exception as e:
        return {
            'error': f'Error generating summary: {str(e)}',
            'summary': 'Unable to generate analysis.',
            'strengths': [],
            'weaknesses': []
        }

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/chat')
def chat():
    return render_template('chat.html')

@app.route('/api/persons', methods=['GET'])
def get_persons():
    """Get list of all saved persons"""
    persons = get_all_saved_persons()
    return jsonify({'success': True, 'persons': persons})

@app.route('/api/load-person', methods=['POST'])
def load_person():
    """Load profile data for a specific person"""
    data = request.json
    person_name = data.get('person_name', '').strip()
    
    if not person_name:
        return jsonify({'error': 'Person name is required'}), 400
    
    profile_data = load_profile_data(person_name)
    
    if not profile_data:
        return jsonify({'error': f'No data found for {person_name}'}), 404
    
    return jsonify({'success': True, 'data': profile_data})

def search_google(query, num_results=5):
    """Search Google and return top result URLs"""
    try:
        # Use DuckDuckGo HTML search as a fallback (no API key needed)
        # Or use a simple Google search via scraping
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        }
        
        # Try Google search
        search_url = f"https://www.google.com/search?q={requests.utils.quote(query)}&num={num_results}"
        response = requests.get(search_url, headers=headers, timeout=10)
        
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            results = []
            
            # Extract search result links
            for link in soup.find_all('a', href=True):
                href = link.get('href', '')
                if href.startswith('/url?q='):
                    # Extract actual URL
                    actual_url = href.split('/url?q=')[1].split('&')[0]
                    if actual_url.startswith('http') and actual_url not in results:
                        results.append(actual_url)
                        if len(results) >= num_results:
                            break
            
            return results
        return []
    except Exception as e:
        print(f"Google search error: {e}")
        return []

def analyze_image_or_video(url):
    """Analyze image or video using Gemini Vision API"""
    try:
        # Use Gemini Vision API
        vision_model = genai.GenerativeModel('gemini-2.5-flash')
        
        response = vision_model.generate_content(
            contents=types.Content(
                parts=[
                    types.Part(
                        file_data=types.FileData(file_uri=url)
                    ),
                    types.Part(text='Please analyze this image or video and provide a detailed description, including any text, objects, people, activities, or relevant information visible.')
                ]
            )
        )
        
        return response.text
    except Exception as e:
        return f"Error analyzing image/video: {str(e)}"

def ai_agent_chat(message, person_data, chat_history):
    """AI Agent with tools to answer questions about a person's profile - supports iterative tool usage"""
    try:
        # Prepare context from person data
        resume_text = person_data.get('resume_text', '')
        analysis = person_data.get('analysis', {})
        scraped_data = person_data.get('scraped_data', {})
        job_description = person_data.get('job_description', '')
        
        # Extract person's name from resume or metadata
        person_name = person_data.get('metadata', {}).get('person_name', '')
        if not person_name:
            # Try to extract from resume text
            name_match = re.search(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', resume_text, re.MULTILINE)
            if name_match:
                person_name = name_match.group(1)
        
        # Extract first name for searching
        first_name = person_name.split()[0] if person_name else ""
        
        # Extract all URLs from resume text
        url_pattern = re.compile(r'https?://[^\s<>"{}|\\^`\[\]]+[^\s<>"{}|\\^`\[\].,;:!?]')
        resume_urls = url_pattern.findall(resume_text)
        
        # Also get URLs from scraped_data (they might have been extracted during initial scraping)
        scraped_urls = []
        for platform, data in scraped_data.items():
            if isinstance(data, dict):
                # Extract URLs from scraped data
                data_str = json.dumps(data)
                scraped_urls.extend(url_pattern.findall(data_str))
        
        # Combine and deduplicate URLs
        all_available_urls = list(set(resume_urls + scraped_urls))
        
        # Map platform names to URLs
        url_map = {}
        for url in all_available_urls:
            url_lower = url.lower()
            if 'github.com' in url_lower:
                url_map['github'] = url
                url_map['github.com'] = url
            elif 'linkedin.com' in url_lower:
                url_map['linkedin'] = url
                url_map['linkedin.com'] = url
            elif 'devpost.com' in url_lower:
                url_map['devpost'] = url
                url_map['devpost.com'] = url
            elif 'kaggle.com' in url_lower:
                url_map['kaggle'] = url
                url_map['kaggle.com'] = url
            elif 'orcid.org' in url_lower:
                url_map['orcid'] = url
                url_map['orcid.org'] = url
        
        # Build initial context
        tools_used = []
        tool_results = []
        max_iterations = 5
        iteration = 0
        
        while iteration < max_iterations:
            iteration += 1
            
            # Build context prompt with accumulated tool results
            available_urls_text = ""
            if all_available_urls:
                available_urls_text = f"""
AVAILABLE LINKS FROM RESUME:
{chr(10).join([f"- {url}" for url in all_available_urls[:20]])}

URL MAPPING (use these when user mentions platforms):
{chr(10).join([f"- {platform}: {url}" for platform, url in url_map.items()])}
"""
            
            # Use full resume text (not truncated) so agent can see all content including projects and links
            # Limit to 8000 chars to avoid token limits, but prioritize showing full content
            resume_text_for_prompt = resume_text[:8000] if len(resume_text) > 8000 else resume_text
            if len(resume_text) > 8000:
                resume_text_for_prompt += "\n\n[Resume text truncated for length, but all key sections should be visible above]"
            
            context_prompt = f"""You are an AI assistant helping to answer questions about a person's professional profile.

PERSON'S NAME: {person_name or 'Unknown'}

FULL RESUME TEXT (with all links, projects, and sections visible):
{resume_text_for_prompt}

PROFILE ANALYSIS:
{json.dumps(analysis, indent=2)[:2000]}

SCRAPED DATA FROM PLATFORMS:
{json.dumps(scraped_data, indent=2)[:3000]}

{available_urls_text}

{f"JOB DESCRIPTION (if relevant): {job_description[:1000]}" if job_description else ""}

{f"TOOL RESULTS FROM PREVIOUS ITERATIONS:\n{chr(10).join(tool_results)}\n" if tool_results else ""}

You have access to 3 tools:
1. **lookup_resume** - Look up information from the FULL RESUME TEXT provided above. The resume text contains:
   - All sections (Projects, Experience, Education, etc.)
   - All links embedded in the resume
   - All project names and descriptions
   - All skills, technologies, and achievements
   Use this tool when you need to reference specific information from the resume. The full resume text is already provided above, so you can answer questions directly from it.
2. **search_website** - Search any website using Firecrawl. Use this when:
   - User asks about a paper, publication, or research
   - User asks to verify information from a specific URL
   - User asks to check a website mentioned in scraped data or resume links
   - User mentions a platform (GitHub, LinkedIn, etc.) - use the URL from AVAILABLE LINKS above
   - You need to search Google first, then scrape the first result
3. **analyze_media** - Analyze images or videos from URLs using Gemini Vision

IMPORTANT INSTRUCTIONS:
- The FULL RESUME TEXT is provided above with ALL content including projects, links, and sections. When user asks about projects or anything in the resume, use lookup_resume tool and reference the full resume text directly.
- ALWAYS check AVAILABLE LINKS FROM RESUME first before searching! If user mentions "GitHub", "LinkedIn", etc., use the actual URL from the resume links above.
- **CRITICAL: Be PROACTIVE with tools!**
  - If user asks "are those projects in his GitHub?" or "is [project] on GitHub?", you MUST:
    1. Get the GitHub URL from AVAILABLE LINKS
    2. Use search_website tool to scrape the GitHub page
    3. Search for the project names in the scraped content
    4. Report if found or not
  - If user asks about a project on DevPost/Kaggle/etc., automatically scrape that platform and search for the project
  - If user asks about videos/images in a project, automatically go to the project page and use analyze_media tool
- If user asks about a paper/publication (e.g., "was he a co-author of paper X?"), you MUST:
  1. First search Google for the paper title
  2. Get the first search result URL
  3. Scrape that URL using search_website tool
  4. Check if the person's name appears in the authors/authorship section
- When using search_website tool, provide the EXACT URL from AVAILABLE LINKS or a valid URL (not platform names like "GitHub")
- You can use MULTIPLE tools in sequence if needed
- Always explain what tool you're using and why
- Be thorough and verify information when asked - don't just say "I don't know", actually scrape and check!
- When user asks about projects, check the FULL RESUME TEXT above - it contains all project information with links

Current conversation history:
{json.dumps(chat_history[-5:], indent=2) if len(chat_history) > 0 else "No previous conversation"}

User's question: {message}

Analyze the question and determine if you need to use any tools. If yes, respond in this JSON format:
{{
    "needs_tool": true/false,
    "tool": "lookup_resume" | "search_website" | "analyze_media" | null,
    "tool_input": "what to search/scrape/analyze",
    "reasoning": "why you need this tool"
}}

If you don't need a tool, respond with:
{{
    "needs_tool": false,
    "final_answer": "your complete answer to the user's question"
}}"""

            # Get AI decision on tool usage
            response = model.generate_content(context_prompt)
            ai_text = response.text.strip()
            
            # Clean JSON response
            if ai_text.startswith('```'):
                ai_text = ai_text.split('```')[1]
                if ai_text.startswith('json'):
                    ai_text = ai_text[4:]
            ai_text = ai_text.strip()
            
            try:
                decision = json.loads(ai_text)
            except:
                # Fallback: try to extract JSON from text
                json_match = re.search(r'\{[^}]+\}', ai_text)
                if json_match:
                    decision = json.loads(json_match.group())
                else:
                    decision = {"needs_tool": False, "final_answer": ai_text}
            
            # If no tool needed, return final answer
            if not decision.get('needs_tool', False):
                final_answer = decision.get('final_answer', ai_text)
                if tool_results:
                    final_answer = "\n\n".join(tool_results) + "\n\n" + final_answer
                return final_answer, ", ".join(tools_used) if tools_used else "lookup_resume"
            
            # Execute tool
            tool_name = decision.get('tool')
            tool_input = decision.get('tool_input', '')
            
            if tool_name == "search_website":
                tools_used.append("search_website")
                
                # First, check if user mentioned a platform - map to actual URL
                tool_input_lower = tool_input.lower() if tool_input else ""
                message_lower = message.lower()
                
                # Check if user is asking about projects on a platform (e.g., "are those in his github?")
                # Extract project names from previous context, chat history, or message
                project_names = []
                if 'project' in message_lower or 'projects' in message_lower or 'those' in message_lower:
                    # Extract from previous tool results
                    for result in tool_results:
                        if 'project' in result.lower():
                            # Extract project names (capitalized words that might be project names)
                            matches = re.findall(r'\b([A-Z][a-zA-Z]+(?:AI|App|System|Platform|Tube)?)\b', result)
                            project_names.extend(matches)
                    
                    # Extract from chat history (previous assistant messages)
                    for chat_msg in chat_history[-3:]:
                        if chat_msg.get('role') == 'assistant':
                            content = chat_msg.get('content', '')
                            # Look for project names in bullet points or numbered lists
                            matches = re.findall(r'(?:^|\n)[\*\-\d+\.]\s*\*?\*?([A-Z][a-zA-Z]+(?:AI|App|System|Platform|Tube)?)', content, re.MULTILINE)
                            project_names.extend(matches)
                    
                    # Extract from resume text - look for project sections
                    resume_lower = resume_text.lower()
                    if 'project' in resume_lower:
                        # Find project names (usually capitalized, might be in bullet points)
                        project_matches = re.findall(r'(?:project|projects?)[\s:]*([A-Z][a-zA-Z\s]+(?:AI|App|System|Platform|Tube)?)', resume_text, re.IGNORECASE)
                        for match in project_matches:
                            # Clean up and extract individual project names
                            cleaned = re.sub(r'\s+', ' ', match.strip())
                            if len(cleaned.split()) <= 3:  # Likely a project name
                                project_names.append(cleaned)
                    
                    # Also check message for specific project mentions
                    if 'inqube' in message_lower:
                        project_names.append('InqubeAI')
                    if 'crunchtube' in message_lower or 'crunch' in message_lower:
                        project_names.append('CrunchTube')
                    
                    # Deduplicate
                    project_names = list(set([p.strip() for p in project_names if len(p.strip()) > 2]))
                
                # Check if it's a platform mention (GitHub, LinkedIn, etc.)
                platform_url = None
                platform_name = None
                for platform, url in url_map.items():
                    if platform in tool_input_lower or platform in message_lower:
                        platform_url = url
                        platform_name = platform
                        break
                
                # If platform URL found, use it and search for projects if mentioned
                if platform_url:
                    tool_results.append(f"🔗 Found {platform_url} from resume links")
                    scraped = scrape_with_firecrawl(platform_url)
                    if scraped:
                        content = scraped.get('markdown', scraped.get('content', ''))[:4000]
                        tool_results.append(f"📄 Scraped {platform_url}:\n{content}")
                        
                        # If user asked about projects, search for them in the scraped content
                        if project_names:
                            content_lower = content.lower()
                            found_projects = []
                            not_found_projects = []
                            for project in project_names:
                                if project.lower() in content_lower:
                                    found_projects.append(project)
                                else:
                                    not_found_projects.append(project)
                            
                            if found_projects:
                                tool_results.append(f"✅ Found projects on {platform_name}: {', '.join(found_projects)}")
                            if not_found_projects:
                                tool_results.append(f"❌ Projects not found on {platform_name}: {', '.join(not_found_projects)}")
                    else:
                        tool_results.append(f"⚠️ Could not scrape {platform_url}")
                    continue
                
                # Check if tool_input is already a valid URL
                if tool_input and (tool_input.startswith('http://') or tool_input.startswith('https://')):
                    scraped = scrape_with_firecrawl(tool_input)
                    if scraped:
                        content = scraped.get('markdown', scraped.get('content', ''))[:2000]
                        tool_results.append(f"📄 Scraped {tool_input}:\n{content}")
                    else:
                        tool_results.append(f"⚠️ Could not scrape {tool_input}")
                    continue
                
                # Check if it's a paper/publication question - search Google first
                if any(keyword in message.lower() for keyword in ['paper', 'publication', 'research', 'article', 'co-author', 'author', 'coauthor']):
                    # Extract paper title from message
                    # Look for quoted text or text after keywords
                    paper_title = None
                    
                    # Try to find quoted text first
                    quoted_match = re.search(r'["\']([^"\']+)["\']', message)
                    if quoted_match:
                        paper_title = quoted_match.group(1)
                    else:
                        # Extract text after keywords like "paper", "publication", etc.
                        for keyword in ['paper', 'publication', 'research', 'article']:
                            pattern = rf'{keyword}\s+(?:called|titled|named|["\']|:)?\s*([^?.,!]+)'
                            match = re.search(pattern, message, re.I)
                            if match:
                                paper_title = match.group(1).strip()
                                break
                    
                    # Fallback: use tool_input or clean the message
                    if not paper_title:
                        paper_title = tool_input if tool_input else message
                        # Remove question words and common phrases
                        paper_title = re.sub(r'^(was|is|are|can|do|does|did|will|would|could|should|tell|check|verify|lookup|look up|search for|find|he|she|they|a|an|the|co-author|coauthor|author of)', '', paper_title, flags=re.I).strip()
                        paper_title = re.sub(r'\?$', '', paper_title).strip()
                    
                    # Search Google for the paper
                    search_query = paper_title
                    google_results = search_google(search_query, num_results=3)
                    
                    if google_results:
                        # Scrape the first result
                        first_url = google_results[0]
                        tool_results.append(f"🔍 Searched Google for '{search_query}' and found: {first_url}")
                        
                        # Scrape the first result
                        scraped = scrape_with_firecrawl(first_url)
                        if scraped:
                            content = scraped.get('markdown', scraped.get('content', ''))[:3000]
                            tool_results.append(f"📄 Scraped content from {first_url}:\n{content}")
                            
                            # Check if person's name is in the content
                            if person_name and person_name.lower() in content.lower():
                                tool_results.append(f"✅ Found person's name '{person_name}' in the content!")
                            elif first_name and first_name.lower() in content.lower():
                                tool_results.append(f"⚠️ Found first name '{first_name}' in the content (partial match)")
                            else:
                                tool_results.append(f"❌ Person's name '{person_name}' not found in the content")
                        else:
                            tool_results.append(f"⚠️ Could not scrape {first_url}")
                    else:
                        tool_results.append(f"⚠️ No Google search results found for '{search_query}'")
                
                elif tool_input:
                    # Check if it's in available URLs
                    matching_url = None
                    for url in all_available_urls:
                        if tool_input.lower() in url.lower() or url.lower() in tool_input.lower():
                            matching_url = url
                            break
                    
                    if matching_url:
                        scraped = scrape_with_firecrawl(matching_url)
                        if scraped:
                            content = scraped.get('markdown', scraped.get('content', ''))[:2000]
                            tool_results.append(f"📄 Scraped {matching_url}:\n{content}")
                        else:
                            tool_results.append(f"⚠️ Could not scrape {matching_url}")
                    else:
                        # Try as direct URL
                        url = tool_input
                        if not url.startswith('http'):
                            url = 'https://' + url
                        scraped = scrape_with_firecrawl(url)
                        if scraped:
                            content = scraped.get('markdown', scraped.get('content', ''))[:2000]
                            tool_results.append(f"📄 Scraped {url}:\n{content}")
                        else:
                            tool_results.append(f"⚠️ Could not scrape {url}. Available URLs from resume: {', '.join(all_available_urls[:5])}")
                else:
                    # No tool_input provided - suggest available URLs
                    tool_results.append(f"⚠️ No URL specified. Available URLs from resume: {', '.join(all_available_urls[:5])}")
                
                # Continue to next iteration to process results
                continue
                
            elif tool_name == "analyze_media":
                tools_used.append("analyze_media")
                url = tool_input
                if url:
                    result = analyze_image_or_video(url)
                    tool_results.append(f"🖼️ Analyzed media at {url}:\n{result}")
                continue
                
            elif tool_name == "lookup_resume":
                tools_used.append("lookup_resume")
                # Already have resume context, just continue
                continue
            else:
                # Unknown tool or no tool needed
                break
        
        # Generate final answer with all tool results
        # Use full resume text for final answer generation
        resume_text_final = resume_text[:6000] if len(resume_text) > 6000 else resume_text
        
        final_prompt = f"""Based on the following information, provide a comprehensive answer to the user's question.

PERSON'S NAME: {person_name or 'Unknown'}

FULL RESUME TEXT (with all projects, links, and sections):
{resume_text_final}

TOOL RESULTS:
{chr(10).join(tool_results) if tool_results else 'No additional tool results'}

USER'S QUESTION: {message}

Provide a clear, accurate answer based on all the information gathered. Reference specific projects, links, and sections from the resume when relevant."""
        
        final_response = model.generate_content(final_prompt)
        answer = final_response.text
        
        return answer, ", ".join(tools_used) if tools_used else "lookup_resume"
        
    except Exception as e:
        return f"Error processing chat: {str(e)}", None

@app.route('/api/chat', methods=['POST'])
def chat_api():
    """Handle chat messages from AI agent"""
    try:
        data = request.json
        message = data.get('message', '').strip()
        person_data = data.get('person_data', {})
        chat_history = data.get('chat_history', [])
        
        if not message:
            return jsonify({'error': 'Message is required'}), 400
        
        if not person_data:
            return jsonify({'error': 'Person data is required'}), 400
        
        # Get AI agent response
        response, tool_used = ai_agent_chat(message, person_data, chat_history)
        
        return jsonify({
            'success': True,
            'response': response,
            'tool_used': tool_used
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/extract-links', methods=['POST'])
def extract_links():
    """Extract links from uploaded CV file"""
    if 'cv' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['cv']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Extract links based on file type
        file_ext = filename.rsplit('.', 1)[1].lower()
        links = extract_links_from_file(file_path, file_ext)
        
        # Clean up uploaded file
        try:
            os.remove(file_path)
        except:
            pass
        
        return jsonify({
            'success': True,
            'links': links,
            'count': len(links)
        })
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'cv' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['cv']
    links = request.form.get('links', '')
    job_description = request.form.get('jobDescription', '')
    person_name = request.form.get('personName', '').strip()
    
    if not person_name:
        return jsonify({'error': 'Person name/ID is required'}), 400
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Extract text from CV
        file_ext = filename.rsplit('.', 1)[1].lower()
        if file_ext == 'pdf':
            cv_text = extract_text_from_pdf(file_path)
        elif file_ext == 'docx':
            cv_text = extract_text_from_docx(file_path)
        else:
            cv_text = extract_text_from_txt(file_path)
        
        # Process links
        link_list = [link.strip() for link in links.split(',') if link.strip()]
        scraped_data = {}
        
        for link in link_list:
            platform_info = extract_platform_info(link)
            if platform_info:
                platform = platform_info['platform']
                identifier = platform_info.get('identifier', '')
                
                if platform == 'github':
                    scraped_data['github'] = scrape_github(identifier)
                elif platform == 'linkedin':
                    scraped_data['linkedin'] = scrape_linkedin(platform_info.get('url', link))
                elif platform == 'devpost':
                    scraped_data['devpost'] = scrape_devpost(identifier)
                elif platform == 'kaggle':
                    scraped_data['kaggle'] = scrape_kaggle(identifier)
                elif platform == 'unknown':
                    # Use AI-powered scraper for unknown websites (ORCID, portfolio, etc.)
                    website_url = platform_info.get('url', link)
                    # Use a cleaner key name based on the domain
                    domain = website_url.split('/')[2] if '/' in website_url else 'website'
                    domain_clean = domain.replace('.', '_').replace('-', '_')
                    scraped_data[domain_clean] = scrape_unknown_website(website_url)
        
        # Generate summary using Gemini (with job description if provided)
        analysis = generate_profile_summary(cv_text, scraped_data, job_description)
        
        # Save all data to filesystem
        saved_dir = save_profile_data(person_name, cv_text, file_path, analysis, scraped_data, job_description)
        
        # Clean up uploaded file (but keep the saved copy)
        try:
            os.remove(file_path)
        except:
            pass
        
        return jsonify({
            'success': True,
            'analysis': analysis,
            'scraped_data': scraped_data,
            'cv_preview': cv_text[:500] + '...' if len(cv_text) > 500 else cv_text,
            'has_job_description': bool(job_description and job_description.strip()),
            'person_name': person_name,
            'saved': saved_dir is not None
        })
    
    return jsonify({'error': 'Invalid file type'}), 400

if __name__ == '__main__':
    app.run(debug=True, port=5000)

